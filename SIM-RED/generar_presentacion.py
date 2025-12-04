#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar documento Word de presentación SIM-RED EXTENDIDO
con diagramas de alta resolución embebidos.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_page_break(doc):
    """Agrega un salto de página"""
    doc.add_page_break()

def add_heading_custom(doc, text, level=1, color=None):
    """Agrega un encabezado con formato personalizado"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_paragraph_formatted(doc, text, bold=False, italic=False, size=None, color=None):
    """Agrega un párrafo con formato personalizado"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def create_presentation_document():
    """Crea el documento Word de presentación"""
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos del documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========== PORTADA ==========
    # Título principal
    title = doc.add_heading('SIM-RED EXTENDIDO', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Sistema de Monitoreo, Análisis y Seguridad para Redes Locales')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    
    # Información del proyecto
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Proyecto Final - Administración de Redes\n')
    run.font.size = Pt(14)
    run = info.add_run('Tecnologías: Bash, AWK, Perl\n')
    run.font.size = Pt(12)
    run = info.add_run('Versión 1.0')
    run.font.size = Pt(12)
    
    add_page_break(doc)
    
    # ========== TABLA DE CONTENIDOS ==========
    add_heading_custom(doc, '📑 Tabla de Contenidos', level=1, color=RGBColor(0, 51, 102))
    
    toc_items = [
        '1. Diagramas del Proyecto',
        '2. Guión de Presentación',
        '3. Demostración Técnica',
        '4. Casos de Uso',
        '5. Banco de Preguntas y Respuestas',
        '6. Evaluación de Completitud',
        '7. Consejos para la Presentación'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    add_page_break(doc)
    
    # ========== DIAGRAMAS DEL PROYECTO ==========
    add_heading_custom(doc, '📊 DIAGRAMAS DEL PROYECTO', level=1, color=RGBColor(0, 51, 102))
    
    # Diagrama 1: Arquitectura del Sistema
    add_heading_custom(doc, '1. Arquitectura del Sistema', level=2, color=RGBColor(0, 102, 204))
    doc.add_paragraph('Este diagrama muestra la estructura completa del sistema SIM-RED EXTENDIDO, incluyendo todos los módulos, bibliotecas, archivos de configuración y flujos de datos.')
    
    # Insertar imagen
    diagram_path = 'diagramas/arquitectura_sistema.png'
    if os.path.exists(diagram_path):
        doc.add_picture(diagram_path, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_page_break(doc)
    
    # Diagrama 2: Flujo de Trabajo
    add_heading_custom(doc, '2. Flujo de Trabajo Principal', level=2, color=RGBColor(0, 102, 204))
    doc.add_paragraph('Este diagrama ilustra el flujo de ejecución del sistema desde el inicio hasta la ejecución de cada módulo y el retorno al menú principal.')
    
    diagram_path = 'diagramas/flujo_trabajo.png'
    if os.path.exists(diagram_path):
        doc.add_picture(diagram_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_page_break(doc)
    
    # Diagrama 3: Estructura de Archivos
    add_heading_custom(doc, '3. Estructura de Archivos', level=2, color=RGBColor(0, 102, 204))
    doc.add_paragraph('Este diagrama presenta la organización de directorios y archivos del proyecto SIM-RED.')
    
    diagram_path = 'diagramas/estructura_archivos.png'
    if os.path.exists(diagram_path):
        doc.add_picture(diagram_path, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_page_break(doc)
    
    # ========== GUIÓN DE PRESENTACIÓN ==========
    add_heading_custom(doc, '🎤 GUIÓN DE PRESENTACIÓN', level=1, color=RGBColor(0, 51, 102))
    
    # Introducción
    add_heading_custom(doc, 'INTRODUCCIÓN (2-3 minutos)', level=2, color=RGBColor(0, 102, 204))
    
    add_heading_custom(doc, 'Saludo y Contexto:', level=3)
    p = doc.add_paragraph()
    p.add_run('"Buenos días/tardes. Hoy les presentaré SIM-RED EXTENDIDO, un sistema completo de monitoreo, análisis y seguridad para redes locales que desarrollé como proyecto final del curso de Administración de Redes."').italic = True
    
    add_heading_custom(doc, 'Problema que Resuelve:', level=3)
    p = doc.add_paragraph()
    p.add_run('"En entornos de red, especialmente en empresas y centros educativos, es fundamental tener control sobre qué dispositivos se conectan, detectar amenazas de seguridad y monitorear el rendimiento. Las soluciones comerciales suelen ser costosas y complejas. SIM-RED ofrece una alternativa gratuita, ligera y efectiva."').italic = True
    
    add_heading_custom(doc, 'Objetivos del Proyecto:', level=3)
    objectives = [
        'Aplicar conocimientos de administración de redes y sistemas',
        'Desarrollar habilidades en Shell scripting, AWK y Perl',
        'Crear una herramienta práctica y funcional',
        'Implementar buenas prácticas de seguridad'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Number')
    
    add_page_break(doc)
    
    # Demostración Técnica
    add_heading_custom(doc, 'DEMOSTRACIÓN TÉCNICA (5-7 minutos)', level=2, color=RGBColor(0, 102, 204))
    
    add_heading_custom(doc, '1. Arquitectura del Sistema', level=3)
    p = doc.add_paragraph()
    p.add_run('"El sistema está construido completamente en Bash, AWK y Perl, sin dependencias de frameworks pesados. Consta de 15 módulos especializados organizados en 5 categorías."').italic = True
    doc.add_paragraph('[Mostrar diagrama de arquitectura]').bold = True
    
    add_heading_custom(doc, '2. Tecnologías Utilizadas', level=3)
    tech_list = [
        ('Bash:', 'Script principal y lógica de negocio'),
        ('AWK (gawk):', 'Procesamiento de datos y análisis estadístico'),
        ('Perl:', 'Generación de informes HTML'),
        ('Herramientas del sistema:', 'arp-scan, nmap, ping, dig')
    ]
    for tech, desc in tech_list:
        p = doc.add_paragraph()
        p.add_run(tech).bold = True
        p.add_run(f' {desc}')
    
    add_heading_custom(doc, '3. Funcionalidades Principales', level=3)
    
    categories = [
        ('Categoría 1: Monitoreo de Dispositivos', 
         'Permite verificar qué dispositivos están conectados, comparándolos con una lista de autorizados, validando horarios permitidos y detectando intrusos.'),
        ('Categoría 2: Análisis de Rendimiento',
         'Mide latencia, tráfico de red y genera estadísticas para diagnóstico de problemas de rendimiento.'),
        ('Categoría 3: Seguridad',
         'Detecta ataques de spoofing, monitorea la tabla ARP, escanea puertos y verifica integridad de archivos de configuración.'),
        ('Categoría 4: Informes y Configuración',
         'Genera informes completos en HTML/TXT, gestiona logs y permite configurar el sistema de forma interactiva.'),
        ('Categoría 5: Sistema',
         'Verifica automáticamente las dependencias y ofrece instalarlas si faltan.')
    ]
    
    for cat, desc in categories:
        p = doc.add_paragraph()
        p.add_run(cat).bold = True
        doc.add_paragraph(f'"{desc}"').italic = True
    
    add_page_break(doc)
    
    # Demostración en Vivo
    add_heading_custom(doc, 'DEMOSTRACIÓN EN VIVO (3-5 minutos)', level=2, color=RGBColor(0, 102, 204))
    
    demo_steps = [
        ('Paso 1: Iniciar el Sistema', 
         'sudo ./sim-red.sh',
         'Al iniciar, el sistema auto-detecta la configuración de red y verifica las herramientas necesarias.'),
        ('Paso 2: Verificar Dispositivos (Opción 1)',
         '',
         'Voy a ejecutar la opción 1 para ver qué dispositivos están conectados en este momento.'),
        ('Paso 3: Generar Informe (Opción 12)',
         '',
         'Ahora generaré un informe completo que incluye todas las verificaciones de seguridad y rendimiento.')
    ]
    
    for step, code, desc in demo_steps:
        add_heading_custom(doc, step, level=3)
        if code:
            p = doc.add_paragraph(code)
            p.style = 'Intense Quote'
        p = doc.add_paragraph()
        p.add_run(f'"{desc}"').italic = True
    
    doc.add_paragraph('[Mostrar el informe HTML generado]').bold = True
    
    add_page_break(doc)
    
    # Casos de Uso
    add_heading_custom(doc, 'CASOS DE USO (2 minutos)', level=2, color=RGBColor(0, 102, 204))
    
    use_cases = [
        ('Caso 1: Red Corporativa', [
            'Control estricto de acceso',
            'Detección de amenazas',
            'Auditorías periódicas'
        ]),
        ('Caso 2: Red Educativa', [
            'Control de horarios',
            'Gestión de ancho de banda',
            'Reportes administrativos'
        ]),
        ('Caso 3: Servidor de Producción', [
            'Monitoreo 24/7',
            'Detección de anomalías',
            'Alta disponibilidad'
        ])
    ]
    
    for case, items in use_cases:
        add_heading_custom(doc, case, level=3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)
    
    # Conclusiones
    add_heading_custom(doc, 'CONCLUSIONES (1-2 minutos)', level=2, color=RGBColor(0, 102, 204))
    
    add_heading_custom(doc, 'Logros del Proyecto:', level=3)
    achievements = [
        'Sistema funcional con 15 módulos especializados',
        'Aplicación práctica de Shell, AWK y Perl',
        'Implementación de buenas prácticas de seguridad',
        'Documentación completa',
        'Código modular y mantenible'
    ]
    for ach in achievements:
        p = doc.add_paragraph()
        p.add_run('✅ ').font.color.rgb = RGBColor(0, 128, 0)
        p.add_run(ach)
    
    add_heading_custom(doc, 'Aprendizajes:', level=3)
    learnings = [
        'Scripting avanzado en Bash',
        'Procesamiento de datos con AWK',
        'Análisis de redes y protocolos',
        'Gestión de logs y reportes'
    ]
    for learn in learnings:
        doc.add_paragraph(learn, style='List Bullet')
    
    add_heading_custom(doc, 'Trabajo Futuro:', level=3)
    future = [
        'Dashboard web en tiempo real',
        'Notificaciones automáticas',
        'Integración con SIEM',
        'API REST'
    ]
    for fut in future:
        doc.add_paragraph(fut, style='List Bullet')
    
    add_page_break(doc)
    
    # ========== BANCO DE PREGUNTAS ==========
    add_heading_custom(doc, '❓ BANCO DE PREGUNTAS Y RESPUESTAS', level=1, color=RGBColor(0, 51, 102))
    
    # Preguntas Generales
    add_heading_custom(doc, 'PREGUNTAS GENERALES', level=2, color=RGBColor(0, 102, 204))
    
    qa_general = [
        ('P1: ¿Qué es SIM-RED EXTENDIDO y para qué sirve?',
         'SIM-RED EXTENDIDO es un sistema de monitoreo, análisis y seguridad para redes locales. Sirve para verificar dispositivos conectados, detectar amenazas de seguridad como spoofing, medir rendimiento de red, y generar informes completos. Es útil en entornos corporativos, educativos o cualquier red que requiera control y monitoreo.'),
        
        ('P2: ¿Por qué desarrollaste este proyecto en Bash/AWK/Perl en lugar de usar Python o un lenguaje moderno?',
         'El objetivo del curso era aplicar conocimientos de administración de sistemas y habilidades en Shell scripting. Bash, AWK y Perl son herramientas nativas de Linux, no requieren instalación adicional, son muy eficientes para tareas de sistema y procesamiento de texto, y son fundamentales para cualquier administrador de sistemas. Además, demuestran dominio de herramientas tradicionales de Unix/Linux.'),
        
        ('P3: ¿Qué problemas resuelve tu proyecto?',
         'Resuelve varios problemas: 1) Control de acceso a la red (dispositivos autorizados), 2) Detección de amenazas (spoofing, dispositivos desconocidos), 3) Monitoreo de rendimiento (latencia, tráfico), 4) Auditoría de seguridad (puertos abiertos, integridad de archivos), 5) Generación de reportes para documentación, 6) Detección de uso no autorizado de VPN/Proxy.')
    ]
    
    for q, a in qa_general:
        p = doc.add_paragraph()
        p.add_run(q).bold = True
        p.add_run('\n\n')
        p.add_run('R: ').bold = True
        p.add_run(a)
        doc.add_paragraph()
    
    add_page_break(doc)
    
    # Preguntas Técnicas - Arquitectura
    add_heading_custom(doc, 'PREGUNTAS TÉCNICAS - ARQUITECTURA', level=2, color=RGBColor(0, 102, 204))
    
    qa_arch = [
        ('P4: ¿Cómo está estructurado el proyecto?',
         'El proyecto sigue una arquitectura modular: sim-red.sh (script principal con menú interactivo), bin/ (15 scripts especializados), lib/ (bibliotecas compartidas), config/ (archivos de configuración), logs/ (registros de actividad), data/ (datos históricos), reports/ (informes generados).'),
        
        ('P5: ¿Por qué separaste las funciones en módulos?',
         'Por varias razones: 1) Mantenibilidad (cada módulo es independiente), 2) Reutilización (bibliotecas comunes evitan duplicación), 3) Escalabilidad (fácil agregar funciones), 4) Debugging (errores aislados), 5) Buenas prácticas (separación de responsabilidades).'),
        
        ('P6: ¿Cómo funciona el sistema de logs?',
         'Cada módulo registra sus actividades en archivos .log específicos con formato [YYYY-MM-DD HH:MM:SS] [NIVEL] Mensaje. Los niveles son: INFO, WARNING, ERROR, ALERT. La opción 13 permite visualizar, filtrar, limpiar y exportar logs. Se retienen según configuración (por defecto 30 días).')
    ]
    
    for q, a in qa_arch:
        p = doc.add_paragraph()
        p.add_run(q).bold = True
        p.add_run('\n\n')
        p.add_run('R: ').bold = True
        p.add_run(a)
        doc.add_paragraph()
    
    add_page_break(doc)
    
    # Preguntas Técnicas - Funcionalidades
    add_heading_custom(doc, 'PREGUNTAS TÉCNICAS - FUNCIONALIDADES', level=2, color=RGBColor(0, 102, 204))
    
    qa_func = [
        ('P7: ¿Cómo funciona la verificación de dispositivos (Opción 1)?',
         'Utiliza arp-scan para escanear la subred y obtener IP, MAC y hostname. Luego: 1) Compara con hosts.conf, 2) Verifica horarios en schedule.conf, 3) Valida que la MAC coincida, 4) Clasifica dispositivos como AUTORIZADO, DESCONOCIDO, FUERA DE HORARIO, MAC NO COINCIDE, 5) Genera resumen y lo registra.'),
        
        ('P8: ¿Cómo detectas ataques de spoofing (Opción 2)?',
         'Lee la tabla ARP (/proc/net/arp) y detecta: 1) IP Spoofing (misma IP con múltiples MACs), 2) MAC Spoofing (misma MAC con múltiples IPs), 3) Cambios históricos (compara con arp_history.dat). Utiliza AWK para procesar y analizar los datos eficientemente.'),
        
        ('P9: ¿Cómo funciona la detección de VPN/Proxy (Opción 3)?',
         'Analiza múltiples indicadores: 1) TTL (detecta cambios vs histórico), 2) Latencia (variaciones inusuales), 3) Puertos VPN (escanea 1194-OpenVPN, 500/4500-IPSec, 1723-PPTP), 4) Probabilidad (calcula score: BAJA/MEDIA/ALTA).'),
        
        ('P10: ¿Cómo mides la latencia (Opciones 4 y 5)?',
         'Opción 4 (Puntual): Hace ping a todos los hosts, extrae estadísticas con AWK, ordena por latencia y genera reporte. Opción 5 (Continua): Mediciones cada segundo, actualiza pantalla en tiempo real, genera gráficas ASCII, alerta cuando se superan umbrales.'),
        
        ('P11: ¿Cómo generas los informes HTML (Opción 12)?',
         'El script generate_report.sh ejecuta múltiples verificaciones y recopila resultados. Luego report_generator.pl (Perl) lee los datos, genera HTML con CSS embebido, incluye tablas y gráficas, aplica colores según severidad, y guarda con timestamp.')
    ]
    
    for q, a in qa_func:
        p = doc.add_paragraph()
        p.add_run(q).bold = True
        p.add_run('\n\n')
        p.add_run('R: ').bold = True
        p.add_run(a)
        doc.add_paragraph()
    
    add_page_break(doc)
    
    # ========== EVALUACIÓN DE COMPLETITUD ==========
    add_heading_custom(doc, '✅ EVALUACIÓN DE COMPLETITUD DEL PROYECTO', level=1, color=RGBColor(0, 51, 102))
    
    add_heading_custom(doc, 'Requisitos Cumplidos', level=2, color=RGBColor(0, 102, 204))
    
    requirements = [
        ('Conocimientos de Administración de Redes', [
            'Escaneo de redes (arp-scan)',
            'Análisis de tabla ARP',
            'Monitoreo de puertos',
            'Gestión de DNS',
            'Medición de latencia y rendimiento',
            'Detección de amenazas de red',
            'Control de acceso basado en horarios'
        ]),
        ('Habilidades de Shell Scripting', [
            'Scripts Bash complejos con funciones',
            'Manejo de argumentos y opciones',
            'Control de flujo (if/case/while/for)',
            'Procesamiento de archivos',
            'Manejo de errores',
            'Códigos de salida',
            'Variables y arrays',
            'Redirección y pipes'
        ]),
        ('Habilidades de AWK', [
            'Procesamiento de archivos delimitados',
            'Cálculos estadísticos (promedio, stddev)',
            'Filtrado y transformación de datos',
            'Generación de reportes formateados',
            'Gráficas ASCII',
            'Análisis de logs'
        ]),
        ('Habilidades de Perl', [
            'Generación de HTML dinámico',
            'Procesamiento de datos',
            'Formateo de reportes',
            'Manipulación de strings'
        ])
    ]
    
    for req, items in requirements:
        add_heading_custom(doc, f'✅ {req}', level=3, color=RGBColor(0, 128, 0))
        for item in items:
            p = doc.add_paragraph()
            p.add_run('✓ ').font.color.rgb = RGBColor(0, 128, 0)
            p.add_run(item)
    
    add_page_break(doc)
    
    # Conclusión de Evaluación
    add_heading_custom(doc, 'Conclusión de Evaluación', level=2, color=RGBColor(0, 102, 204))
    
    p = doc.add_paragraph()
    p.add_run('El proyecto está COMPLETO y CUMPLE AMPLIAMENTE con los objetivos:').bold = True
    
    conclusions = [
        'Demuestra dominio de administración de redes',
        'Demuestra habilidades avanzadas en Shell, AWK y Perl',
        'Es funcional y útil en entornos reales',
        'Está bien documentado',
        'Sigue buenas prácticas de programación'
    ]
    
    for conc in conclusions:
        p = doc.add_paragraph()
        p.add_run('✅ ').font.color.rgb = RGBColor(0, 128, 0)
        p.add_run(conc)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Calificación estimada: 95-100/100').bold = True
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    add_page_break(doc)
    
    # ========== CONSEJOS PARA LA PRESENTACIÓN ==========
    add_heading_custom(doc, '🎯 CONSEJOS PARA LA PRESENTACIÓN', level=1, color=RGBColor(0, 51, 102))
    
    tips_sections = [
        ('Antes de Presentar', [
            'Prueba todas las funciones para asegurarte que funcionan',
            'Ten el sistema ejecutándose en una VM o red de prueba',
            'Prepara ejemplos de informes generados',
            'Revisa los logs para mostrar ejemplos reales',
            'Ten a mano el código de 2-3 scripts para mostrar si preguntan'
        ]),
        ('Durante la Presentación', [
            'Habla con confianza, conoces tu proyecto',
            'Usa los diagramas para explicar la arquitectura',
            'Haz una demo en vivo (aunque sea breve)',
            'Muestra el código si preguntan, pero no te pierdas en detalles',
            'Controla el tiempo (10-15 minutos típicamente)'
        ]),
        ('Al Responder Preguntas', [
            'Escucha la pregunta completa antes de responder',
            'Si no sabes algo, sé honesto pero sugiere cómo lo investigarías',
            'Relaciona tus respuestas con conceptos del curso',
            'Usa ejemplos concretos de tu código',
            'Mantén la calma, es TU proyecto, tú eres el experto'
        ])
    ]
    
    for section, tips in tips_sections:
        add_heading_custom(doc, section, level=2, color=RGBColor(0, 102, 204))
        for tip in tips:
            p = doc.add_paragraph()
            p.add_run('✅ ').font.color.rgb = RGBColor(0, 128, 0)
            p.add_run(tip)
    
    add_heading_custom(doc, 'Frases Útiles', level=2, color=RGBColor(0, 102, 204))
    phrases = [
        '"Esa es una excelente pregunta. En mi implementación..."',
        '"Consideré esa opción, pero elegí X porque..."',
        '"Eso sería una mejora futura interesante..."',
        '"Déjame mostrarte el código específico para eso..."',
        '"Basándome en lo que aprendimos en clase sobre..."'
    ]
    for phrase in phrases:
        p = doc.add_paragraph()
        p.add_run('• ').font.color.rgb = RGBColor(0, 102, 204)
        p.add_run(phrase).italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Mensaje final
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final.add_run('¡Éxito en tu presentación! 🚀')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.bold = True
    
    # Guardar documento
    output_path = 'presentacion_simred.docx'
    doc.save(output_path)
    print(f"✅ Documento Word creado exitosamente: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        create_presentation_document()
    except Exception as e:
        print(f"❌ Error al crear el documento: {e}")
        import traceback
        traceback.print_exc()
